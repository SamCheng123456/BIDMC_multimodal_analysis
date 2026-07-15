from __future__ import annotations

from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
INPUT_DOCX = ROOT / "docs" / "course_paper_full_dataset_enhanced_gui_cn.docx"
OUTPUT_DOCX = ROOT / "docs" / "course_paper_full_dataset_enhanced_gui_cn_formatted.docx"

SIMSUN = "宋体"
KAITI = "楷体"
TNR = "Times New Roman"


def set_run_font(run, font_name: str = SIMSUN, size_pt: float = 14, bold: bool | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = rpr._add_rFonts()
    ascii_font = TNR if font_name == SIMSUN else font_name
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)


def format_paragraph(
    para,
    font_name: str = SIMSUN,
    size_pt: float = 14,
    bold: bool = False,
    align=None,
    line_spacing: float = 1.5,
    first_line_indent: bool = False,
) -> None:
    if align is not None:
        para.alignment = align
    para.paragraph_format.line_spacing = line_spacing
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.first_line_indent = Pt(size_pt * 2) if first_line_indent else None
    for run in para.runs:
        set_run_font(run, font_name, size_pt, bold)


def format_style(doc: Document, style_name: str, font_name: str, size_pt: float, bold: bool, line_spacing: float) -> None:
    style = doc.styles[style_name]
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = rpr._add_rFonts()
    ascii_font = TNR if font_name == SIMSUN else font_name
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    style.paragraph_format.line_spacing = line_spacing
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)


def set_para_text(para, text: str) -> None:
    para.text = text


def replace_text_runs(para, replacements: dict[str, str]) -> None:
    text = para.text
    new_text = text
    for old, new in replacements.items():
        new_text = new_text.replace(old, new)
    if new_text != text:
        para.text = new_text


def move_paragraph_after(paragraph_to_move, anchor_paragraph) -> None:
    element = paragraph_to_move._p
    anchor = anchor_paragraph._p
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    anchor.addnext(element)


def main() -> None:
    doc = Document(INPUT_DOCX)

    format_style(doc, "Normal", SIMSUN, 14, False, 1.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        format_style(doc, name, SIMSUN, 14, True, 1.5)
    format_style(doc, "Caption", SIMSUN, 10.5, True, 1.0)

    if len(doc.paragraphs) > 0:
        format_paragraph(doc.paragraphs[0], SIMSUN, 18, True, WD_ALIGN_PARAGRAPH.CENTER, 1.5)
    if len(doc.paragraphs) > 2:
        format_paragraph(doc.paragraphs[2], SIMSUN, 18, True, WD_ALIGN_PARAGRAPH.CENTER, 1.5)
    if len(doc.paragraphs) > 3:
        set_para_text(doc.paragraphs[3], "程一鸣  20243008391")
        format_paragraph(doc.paragraphs[3], KAITI, 16, False, WD_ALIGN_PARAGRAPH.CENTER, 1.5)
    if len(doc.paragraphs) > 4:
        set_para_text(doc.paragraphs[4], "")

    heading_map = {
        "1 引言": "一、引言",
        "2 数据来源与实验方法": "二、数据来源与实验方法",
        "3 数据预处理": "三、数据预处理",
        "4 数据分析与特征提取": "四、数据分析与特征提取",
        "5 可视化结果": "五、可视化结果",
        "6 多模态融合与扩展分析": "六、多模态融合与扩展分析",
        "7 实时可视化界面设计": "七、实时可视化界面设计",
        "7.1 全数据集扩展结果": "（一）全数据集扩展结果",
        "8 分析与讨论": "八、分析与讨论",
        "9 结论": "九、结论",
        "10 参考文献": "十、参考文献",
        "11 附录（代码说明）": "十一、附录（代码说明）",
    }
    for para in doc.paragraphs:
        text = para.text.strip()
        if text in heading_map:
            para.text = heading_map[text]
            para.style = doc.styles["Heading 2" if text.startswith("7.1") else "Heading 1"]

    reference_replacements = {
        "从图 1 到图 3": "从图 2 到图 4",
        "图 4 至图 6": "图 5 至图 7",
        "图 7 的频谱": "图 8 的频谱",
        "图 8 的 STFT": "图 9 的 STFT",
        "图 9 保留记录": "图 10 保留记录",
        "图 18 至图 23": "图 17 至图 22",
        "图 19 的误差排序": "图 18 的误差排序",
        "图 21 和图 22 则从监护仪 Numerics 角度验证": "图 20 和图 21 则从监护仪 Numerics 角度验证",
        "从图 17 可以看出": "从图 15 可以看出",
    }
    for para in doc.paragraphs:
        replace_text_runs(para, reference_replacements)

    fig_caption_indices = [28, 34, 36, 38, 44, 46, 48, 50, 52, 60, 69, 71, 73, 75, 81, 82, 86, 88, 90, 92, 94, 96]
    fig_image_indices = [29, 33, 35, 37, 43, 45, 47, 49, 51, 59, 67, 70, 72, 74, 80, 83, 87, 89, 91, 93, 95, 97]
    fig_captions = [
        "图 1 ECG、PPG 与呼吸信号全数据集多模态分析流程",
        "图 2 记录 01 的原始 ECG、PPG 与呼吸信号",
        "图 3 记录 01 的滤波后 ECG、PPG 与呼吸信号",
        "图 4 标准化后三路信号叠加比较",
        "图 5 ECG R 峰检测结果",
        "图 6 PPG 脉搏峰检测结果",
        "图 7 呼吸峰检测结果",
        "图 8 ECG、PPG 与 RESP 的归一化频谱",
        "图 9 ECG、PPG 与 RESP 的 STFT 时频图",
        "图 10 记录 01 分窗心率、脉率和呼吸率趋势",
        "图 11 ECG-PPG 峰值时延分布",
        "图 12 ECG 心率与 PPG 脉率 Bland-Altman 一致性分析",
        "图 13 心率与呼吸率的心肺耦合分析",
        "图 14 多模态融合特征图",
        "图 15 实时多模态信号可视化界面",
        "图 16 实时回放界面的功能结构",
        "图 17 全数据集生命体征分布",
        "图 18 全部记录 ECG 心率与 PPG 脉率误差排序",
        "图 19 全数据集多模态特征相关性矩阵",
        "图 20 信号估计值与监护仪 Numerics 的误差分布",
        "图 21 信号估计值与监护仪 Numerics 的一致性散点图",
        "图 22 记录级 ECG 与监护仪 HR 误差排序及质量标记",
    ]
    original_paragraphs = list(doc.paragraphs)
    fig_caption_paragraphs = []
    fig_image_paragraphs = []
    for idx, image_idx, caption in zip(fig_caption_indices, fig_image_indices, fig_captions):
        if idx < len(doc.paragraphs):
            para = doc.paragraphs[idx]
            para.text = caption
            para.style = doc.styles["Caption"]
            fig_caption_paragraphs.append(original_paragraphs[idx])
            fig_image_paragraphs.append(original_paragraphs[image_idx])

    table_captions = {
        25: "表 1 全部 53 条 BIDMC 记录的主要生命体征与 HRV 特征统计",
        65: "表 2 全数据集多模态融合与质量控制摘要",
    }
    for idx, caption in table_captions.items():
        if idx < len(doc.paragraphs):
            para = doc.paragraphs[idx]
            para.text = caption
            para.style = doc.styles["Caption"]

    caption_index_set = set(fig_caption_indices) | set(table_captions)
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if i == 2:
            format_paragraph(para, SIMSUN, 18, True, WD_ALIGN_PARAGRAPH.CENTER, 1.5)
        elif i == 3:
            format_paragraph(para, KAITI, 16, False, WD_ALIGN_PARAGRAPH.CENTER, 1.5)
        elif i in caption_index_set:
            para.style = doc.styles["Caption"]
            format_paragraph(para, SIMSUN, 10.5, True, WD_ALIGN_PARAGRAPH.CENTER, 1.0)
        elif para.style.name in ("Heading 1", "Heading 2", "Heading 3"):
            format_paragraph(para, SIMSUN, 14, True, WD_ALIGN_PARAGRAPH.LEFT, 1.5)
        elif text:
            format_paragraph(para, SIMSUN, 14, False, None, 1.5, first_line_indent=(i >= 16))

    # Figure captions belong below figures. Some generated document revisions placed
    # captions above the image, which can separate the visual and its title at a page break.
    for caption_para, image_para in zip(fig_caption_paragraphs, fig_image_paragraphs):
        move_paragraph_after(caption_para, image_para)

    for table in doc.tables:
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing = 1.0
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    for run in para.runs:
                        set_run_font(run, SIMSUN, 10.5, False)

    tmp_docx = OUTPUT_DOCX.with_suffix(".tmp.docx")
    doc.save(tmp_docx)

    media_replacements = {
        "word/media/image1.png": ROOT / "figures/full_dataset/fig0_overall_analysis_workflow.png",
        "word/media/image2.png": ROOT / "figures/01/fig1_raw_multimodal_signals.png",
        "word/media/image3.png": ROOT / "figures/01/fig2_filtered_multimodal_signals.png",
        "word/media/image4.png": ROOT / "figures/01/fig3_standardized_overlay.png",
        "word/media/image5.png": ROOT / "figures/01/fig4_ecg_peaks.png",
        "word/media/image6.png": ROOT / "figures/01/fig5_ppg_peaks.png",
        "word/media/image7.png": ROOT / "figures/01/fig6_resp_peaks.png",
        "word/media/image8.png": ROOT / "figures/01/fig7_frequency_spectra.png",
        "word/media/image9.png": ROOT / "figures/01/fig8_stft_time_frequency.png",
        "word/media/image10.png": ROOT / "figures/01/fig9_windowed_vital_rates.png",
        "word/media/image11.png": ROOT / "figures/advanced_fusion/fig14_pulse_arrival_time_distribution.png",
        "word/media/image12.png": ROOT / "figures/advanced_fusion/fig15_bland_altman_hr_pr.png",
        "word/media/image13.png": ROOT / "figures/advanced_fusion/fig16_cardiorespiratory_coupling.png",
        "word/media/image14.png": ROOT / "figures/advanced_fusion/fig17_multimodal_fusion_feature_map.png",
        "word/media/image15.png": ROOT / "figures/gui/fig18_gui_preview.png",
        "word/media/image16.png": ROOT / "figures/gui/fig_gui_software_explainer.png",
        "word/media/image17.png": ROOT / "figures/full_dataset/fig19_full_dataset_vital_sign_distributions.png",
        "word/media/image18.png": ROOT / "figures/full_dataset/fig20_ranked_hr_pr_error_all_records.png",
        "word/media/image19.png": ROOT / "figures/full_dataset/fig21_full_dataset_feature_correlation.png",
        "word/media/image20.png": ROOT / "figures/full_dataset/fig23_signal_vs_monitor_error_distribution.png",
        "word/media/image21.png": ROOT / "figures/full_dataset/fig24_signal_monitor_agreement_scatter.png",
        "word/media/image22.png": ROOT / "figures/full_dataset/fig25_record_quality_flags.png",
    }

    zip_tmp = OUTPUT_DOCX.with_suffix(".ziptmp.docx")
    with ZipFile(tmp_docx, "r") as zin, ZipFile(zip_tmp, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename in media_replacements:
                data = media_replacements[item.filename].read_bytes()
            else:
                data = zin.read(item.filename)
            zout.writestr(item, data)
    shutil.move(zip_tmp, OUTPUT_DOCX)
    tmp_docx.unlink(missing_ok=True)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
