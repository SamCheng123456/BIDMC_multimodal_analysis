from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


TITLE = "基于 ECG、PPG 与呼吸信号的多模态生命体征特征提取与分析"


def project_root() -> Path:
    return Path(__file__).resolve().parents[1]


def set_east_asian_font(run, font_name: str = "宋体") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_east_asian_font(run)
    run.font.size = Pt(9)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "BFBFBF")


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 14, "1F4E79"),
        ("Heading 3", 12, "1F4E79"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("《劳动教育实践 4——专业劳动实践》\n期末课程论文")
    set_east_asian_font(run, "黑体")
    run.font.size = Pt(18)
    run.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    set_east_asian_font(run, "黑体")
    run.font.size = Pt(22)
    run.bold = True

    for _ in range(5):
        doc.add_paragraph()

    info = [
        "课程名称：劳动教育实践 4——专业劳动实践",
        "论文题目：基于 ECG、PPG 与呼吸信号的多模态生命体征特征提取与分析",
        "姓名：请填写",
        "学号：请填写",
        "班级：请填写",
        "提交日期：2026 年 7 月",
    ]
    for item in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        set_east_asian_font(run, "宋体")
        run.font.size = Pt(14)

    doc.add_section(WD_SECTION_START.NEW_PAGE)


def add_paragraph(doc: Document, text: str, first_line: bool = True) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(6)
    if first_line:
        p.paragraph_format.first_line_indent = Pt(22)
    run = p.add_run(text)
    set_east_asian_font(run)
    run.font.size = Pt(11)


def add_caption(doc: Document, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(caption)
    set_east_asian_font(run, "宋体")
    run.font.size = Pt(9)


def add_figure(doc: Document, path: Path, caption: str, width_cm: float = 14.5) -> None:
    if not path.exists():
        add_paragraph(doc, f"[缺失图片：{path}]", first_line=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_dataframe_table(doc: Document, df: pd.DataFrame, caption: str, headers: list[str], rows: list[list[str]]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    set_east_asian_font(run, "黑体")
    run.font.size = Pt(10)

    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = True
    set_table_borders(table)
    for i, header in enumerate(headers):
        shade_cell(table.rows[0].cells[i], "E8EEF5")
        set_cell_text(table.rows[0].cells[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def fmt(value: float, digits: int = 2) -> str:
    return f"{value:.{digits}f}"


def add_paper_body(doc: Document, root: Path) -> None:
    summary = pd.read_csv(root / "results" / "all_records_summary.csv", dtype={"record_id": str})
    fusion = pd.read_csv(root / "results" / "advanced_fusion_summary.csv", dtype={"record_id": str})
    summary["record_id"] = summary["record_id"].str.zfill(2)
    fusion["record_id"] = fusion["record_id"].str.zfill(2)

    doc.add_heading("摘要", level=1)
    add_paragraph(
        doc,
        "为提高对生命体征状态的综合理解，本文基于 PhysioNet 平台公开的 BIDMC PPG and Respiration Dataset，"
        "选取 01、02、03 三条代表性记录，对同步采集的 ECG、PPG 与呼吸信号进行多模态分析。研究首先对三路信号进行带通滤波和标准化处理，"
        "其中 ECG 采用 0.5-40 Hz 带通滤波，PPG 采用 0.5-8 Hz 带通滤波，呼吸信号采用 0.05-1 Hz 带通滤波。随后从时域、频域和时频域三个角度开展分析，"
        "提取心率、脉率、呼吸率、SDNN、RMSSD、RMS、主频率等特征，并进一步构建 ECG-PPG 一致性分析、ECG-PPG 峰值时延估计和心肺耦合分析。"
        "结果显示，记录 01 和 02 中 ECG 心率与 PPG 脉率较为接近，平均绝对误差分别为 0.59 bpm 和 0.66 bpm；记录 03 的差异较大，平均绝对误差为 3.40 bpm，"
        "提示 PPG 波形质量和外周循环因素可能影响脉率估计。本文还设计了一个轻量级实时可视化界面，用于同步回放 ECG、PPG 和呼吸信号。研究表明，"
        "多模态生理信号联合分析能够从心电活动、外周脉搏和呼吸节律三个层面刻画生命体征状态，较单一信号分析具有更完整的解释能力。"
    )
    add_paragraph(doc, "关键词：ECG；PPG；呼吸信号；BIDMC；多模态融合；生命体征；时频分析", first_line=False)

    doc.add_heading("1 引言", level=1)
    for text in [
        "生物医学信号处理是生物医学工程专业的重要基础内容，其目标是从复杂、含噪的生理信号中提取具有医学意义的信息。"
        "在临床监护、可穿戴健康设备和远程医疗场景中，心电、光电容积脉搏波和呼吸信号是最常见的生命体征信号。"
        "ECG 直接反映心脏电活动，能够较准确地定位心搏事件；PPG 通过光学方式反映外周组织血容量变化，常用于脉率和血氧相关监测；"
        "呼吸信号则反映通气节律，是评估患者基础生理状态的重要指标。三类信号从不同生理机制出发描述人体状态，具有天然互补性。",
        "传统课程实验往往针对单一信号进行滤波、频谱或峰值检测，但在真实监护场景中，单一路信号容易受到噪声、运动、传感器接触不良或局部生理变化影响。"
        "例如，PPG 信号虽然易于采集，但其波形会受到外周灌注、探头位置和运动伪迹影响；呼吸信号频率较低，容易受到基线漂移和体动干扰。"
        "因此，将 ECG、PPG 和呼吸信号进行联合分析，有助于从多个角度相互验证生命体征估计结果，并提高结果解释的可靠性。",
        "本文围绕“生物医学信号处理公开数据集独立分析”的课程任务，选择 BIDMC PPG and Respiration Dataset 作为研究对象。"
        "本文的主要工作包括：第一，完成公开数据集获取、读取和项目化管理；第二，对 ECG、PPG、RESP 三路信号进行针对性预处理；"
        "第三，综合使用时域、频域和时频分析方法提取特征；第四，利用跨记录统计、相关性分析、Bland-Altman 一致性分析和 ECG-PPG 峰值时延估计体现多模态融合思想；"
        "第五，设计实时可视化 GUI，为信号动态观察提供工程扩展。"
    ]:
        add_paragraph(doc, text)

    doc.add_heading("2 数据来源与实验方法", level=1)
    for text in [
        "本文使用的数据集为 PhysioNet 平台发布的 BIDMC PPG and Respiration Dataset v1.0.0。该数据集来源于 Beth Israel Deaconess Medical Center 的危重患者监护记录，"
        "包含 ECG、PPG/PLETH、阻抗呼吸信号、生命体征数值以及呼吸相关标注。数据集总共包含 53 条记录，每条记录约 8 min，波形信号采样率为 125 Hz。"
        "由于课程论文更强调完整处理流程和独立分析，本文选取 bidmc_01、bidmc_02 和 bidmc_03 三条记录作为代表性样本，每条记录按 30 s 窗口切分，共得到 60 个分析窗口。",
        "实验流程包括数据读取、信号预处理、特征提取、可视化展示和多模态融合分析。原始 CSV 文件包含 Time、RESP、PLETH、V、AVR 和 II 等列，本文选用 II 导联作为 ECG 信号，"
        "PLETH 作为 PPG 信号，RESP 作为呼吸信号。读取后首先计算采样率并检查时间间隔，然后分别按信号类型设置滤波参数。"
        "滤波完成后保存处理后数据，以便后续复现实验和论文作图。"
    ]:
        add_paragraph(doc, text)

    rows = []
    for _, row in summary.iterrows():
        rows.append(
            [
                row["record_id"],
                fmt(row["ecg_hr_mean"]) + "±" + fmt(row["ecg_hr_std"]),
                fmt(row["ppg_pr_mean"]) + "±" + fmt(row["ppg_pr_std"]),
                fmt(row["resp_rr_mean"]) + "±" + fmt(row["resp_rr_std"]),
                fmt(row["ecg_sdnn_mean"]),
                fmt(row["ecg_rmssd_mean"]),
                fmt(row["hr_pr_error_mean"]),
            ]
        )
    add_dataframe_table(
        doc,
        summary,
        "表 1 三条 BIDMC 记录的主要生命体征与 HRV 特征",
        ["记录", "ECG心率(bpm)", "PPG脉率(bpm)", "呼吸率(次/min)", "SDNN(ms)", "RMSSD(ms)", "HR-PR误差(bpm)"],
        rows,
    )

    doc.add_heading("3 数据预处理", level=1)
    for text in [
        "生理信号通常包含基线漂移、高频噪声和局部伪迹。为了保证后续特征提取的稳定性，本文对三类信号分别设置滤波参数。"
        "ECG 信号的主要诊断频段集中在较低至中频范围，本文使用 0.5-40 Hz 四阶 Butterworth 带通滤波器，用于削弱基线漂移和高频噪声。"
        "PPG 信号的脉搏节律主要位于心率相关频段，本文使用 0.5-8 Hz 带通滤波器保留脉搏波主体。呼吸信号变化较慢，本文使用 0.05-1 Hz 带通滤波器突出呼吸周期。",
        "滤波器实现采用二阶节形式进行零相位滤波。与普通分子分母系数形式相比，二阶节形式在低截止频率场景下具有更好的数值稳定性，能够避免滤波结果出现异常放大。"
        "滤波后，本文对三路信号进行零均值单位方差标准化，使 ECG、PPG 和 RESP 可以在同一图中叠加比较。标准化只用于可视化和跨模态比较，不改变原始物理信号的相对时间关系。"
    ]:
        add_paragraph(doc, text)

    fig_root = root / "figures"
    add_figure(doc, fig_root / "01" / "fig1_raw_multimodal_signals.png", "图 1 记录 01 的原始 ECG、PPG 与呼吸信号")
    add_figure(doc, fig_root / "01" / "fig2_filtered_multimodal_signals.png", "图 2 记录 01 的滤波后 ECG、PPG 与呼吸信号")
    add_figure(doc, fig_root / "01" / "fig3_standardized_overlay.png", "图 3 标准化后三路信号叠加比较", width_cm=14.0)

    doc.add_heading("4 数据分析与特征提取", level=1)
    for text in [
        "本文至少使用三类分析方法。第一类为时域分析，主要通过峰值检测定位 ECG R 峰、PPG 脉搏峰和呼吸峰，并由相邻峰间期计算心率、脉率和呼吸率。"
        "ECG R 峰检测设置最小峰距约 0.35 s，以避免同一次心搏被重复识别；PPG 峰值检测设置最小峰距约 0.40 s；呼吸峰值检测设置更长的最小峰距，以适应较低的呼吸频率。"
        "在 ECG 特征中，本文进一步计算 SDNN 和 RMSSD。SDNN 反映 RR 间期总体波动，RMSSD 反映相邻 RR 间期短期变化，是常用 HRV 指标。",
        "第二类为频域分析。本文对滤波后信号进行快速傅里叶变换，计算归一化频谱，并在相应频率范围内寻找主频率。ECG 和 PPG 的主频率与心率或脉率相关，"
        "呼吸信号的主频率则对应呼吸节律。频域分析能够验证峰值检测得到的生命体征估计是否落在合理范围内。",
        "第三类为时频分析。本文采用短时傅里叶变换，将信号分解为时间-频率二维表示，用于观察频率成分是否随时间变化。"
        "对于监护信号而言，生命体征并非完全恒定，因此 STFT 图可直观展示心率、脉率和呼吸频率在不同时间段的能量分布。"
    ]:
        add_paragraph(doc, text)

    add_figure(doc, fig_root / "01" / "fig4_ecg_peaks.png", "图 4 ECG R 峰检测结果", width_cm=14.0)
    add_figure(doc, fig_root / "01" / "fig5_ppg_peaks.png", "图 5 PPG 脉搏峰检测结果", width_cm=14.0)
    add_figure(doc, fig_root / "01" / "fig6_resp_peaks.png", "图 6 呼吸峰检测结果", width_cm=14.0)
    add_figure(doc, fig_root / "01" / "fig7_frequency_spectra.png", "图 7 ECG、PPG 与 RESP 的归一化频谱", width_cm=14.0)
    add_figure(doc, fig_root / "01" / "fig8_stft_time_frequency.png", "图 8 ECG、PPG 与 RESP 的 STFT 时频图")

    doc.add_heading("5 可视化结果", level=1)
    for text in [
        "从图 1 到图 3 可以看出，三类信号在时间尺度和波形形态上具有明显差异。ECG 中尖锐的 R 峰对应心脏电活动，PPG 波形呈现相对平滑的脉搏波周期，"
        "呼吸信号则表现为低频缓慢起伏。滤波后，基线漂移和部分高频噪声得到抑制，三路信号周期结构更加清楚。",
        "图 4 至图 6 展示了时域峰值检测结果。ECG R 峰较尖锐，检测结果较稳定；PPG 脉搏峰受波形宽度和局部噪声影响更明显；呼吸峰间隔较长，"
        "其频率明显低于心率和脉率。图 7 进一步从频域角度展示了这种差异，ECG 和 PPG 的主要能量集中在心率相关频段，而 RESP 主频位于更低频率范围。",
        "图 8 的 STFT 结果说明三类信号的频率成分随时间存在轻微变化。与静态频谱相比，时频图更适合观察监护信号在不同时间段的状态变化。"
        "在课程论文要求中，时频分析不仅满足方法多样性要求，也有助于解释真实生理信号的非平稳性。"
    ]:
        add_paragraph(doc, text)

    add_figure(doc, fig_root / "01" / "fig9_windowed_vital_rates.png", "图 9 记录 01 分窗心率、脉率和呼吸率趋势", width_cm=14.0)
    add_figure(doc, fig_root / "all_records" / "fig11_cross_record_vital_signs.png", "图 10 三条记录的生命体征跨记录比较", width_cm=14.0)
    add_figure(doc, fig_root / "all_records" / "fig12_hr_pr_error_comparison.png", "图 11 ECG 心率与 PPG 脉率平均绝对误差比较", width_cm=13.0)
    add_figure(doc, fig_root / "all_records" / "fig13_cross_record_feature_correlation.png", "图 12 跨记录多模态特征相关性矩阵", width_cm=12.5)

    doc.add_heading("6 多模态融合与扩展分析", level=1)
    for text in [
        "为体现多模态融合思想，本文不只对三路信号分别提取特征，还进一步分析不同模态之间的一致性和互补性。"
        "首先，本文比较 ECG 心率与 PPG 脉率。记录 01 和 02 中二者平均绝对误差分别为 0.59 bpm 和 0.66 bpm，说明在这两条记录中 PPG 脉率能够较好反映 ECG 心率变化。"
        "记录 03 中平均绝对误差增大到 3.40 bpm，提示 PPG 波形质量、外周循环状态或峰值检测稳定性可能影响脉率估计。",
        "其次，本文采用 Bland-Altman 图分析 ECG 心率与 PPG 脉率的一致性。与简单相关系数不同，Bland-Altman 图直接观察两种测量方法之间的差值和一致性范围，"
        "更适合判断两种生命体征估计方式是否可以相互替代。结果显示，记录 01 和 02 的多数点接近零差值附近，而记录 03 的若干窗口差异较大，说明跨记录波形质量差异会影响一致性。",
        "此外，本文估计 ECG R 峰与 PPG 脉搏峰之间的峰值时延。该指标从信号处理角度描述心电活动与外周脉搏波之间的时间关系，但由于本文未进行严格的传感器位置校准和临床血压建模，"
        "因此仅将其作为 ECG-PPG 时间耦合特征，而不直接等同于临床标准脉搏传导时间。最后，本文分析呼吸率与心率之间的关系，用于观察心肺节律之间的统计联系。"
    ]:
        add_paragraph(doc, text)

    fusion_rows = []
    for _, row in fusion.iterrows():
        fusion_rows.append(
            [
                row["record_id"],
                str(int(row["pat_count"])),
                fmt(row["pat_mean_ms"]),
                fmt(row["pat_median_ms"]),
                fmt(row["ecg_ppg_rate_corr"], 3),
                fmt(row["ecg_resp_rate_corr"], 3),
                fmt(row["ppg_resp_rate_corr"], 3),
            ]
        )
    add_dataframe_table(
        doc,
        fusion,
        "表 2 多模态融合扩展特征",
        ["记录", "时延样本数", "平均时延(ms)", "中位时延(ms)", "HR-PR相关", "HR-RR相关", "PR-RR相关"],
        fusion_rows,
    )

    add_figure(doc, fig_root / "advanced_fusion" / "fig14_pulse_arrival_time_distribution.png", "图 13 ECG-PPG 峰值时延分布", width_cm=13.0)
    add_figure(doc, fig_root / "advanced_fusion" / "fig15_bland_altman_hr_pr.png", "图 14 ECG 心率与 PPG 脉率 Bland-Altman 一致性分析", width_cm=14.0)
    add_figure(doc, fig_root / "advanced_fusion" / "fig16_cardiorespiratory_coupling.png", "图 15 心率与呼吸率的心肺耦合分析", width_cm=13.0)
    add_figure(doc, fig_root / "advanced_fusion" / "fig17_multimodal_fusion_feature_map.png", "图 16 多模态融合特征图", width_cm=13.0)

    doc.add_heading("7 实时可视化界面设计", level=1)
    for text in [
        "除离线分析外，本文设计了一个轻量级实时可视化界面，用于回放处理后的 ECG、PPG 和 RESP 三路信号。该界面基于 Python 标准库 Tkinter 和 Matplotlib 实现，"
        "不需要额外复杂框架。用户可以选择记录编号，设置信号显示窗口长度和播放步长，通过播放、暂停、重置和滑动条查看不同时间段的多模态信号。",
        "该 GUI 的作用并不是替代离线统计分析，而是提供一种更直观的工程展示方式。对于多模态信号而言，同步观察三路波形随时间变化有助于理解 ECG R 峰、PPG 脉搏波和呼吸节律之间的相对关系。"
        "在课程任务的鼓励项中，实时可视化界面属于工程扩展内容，可体现项目不仅完成了算法处理，也考虑了结果展示与交互观察。"
    ]:
        add_paragraph(doc, text)

    doc.add_heading("8 分析与讨论", level=1)
    for text in [
        "从三条记录的结果看，bidmc_01 和 bidmc_02 的 ECG 心率均约为 73 bpm，PPG 脉率也接近该范围，说明在这两条记录中两类心血管相关信号具有较高一致性。"
        "bidmc_03 的 ECG 心率约为 64.46 bpm，而 PPG 脉率约为 61.08 bpm，二者差异相对更大，同时该记录的 SDNN 和 RMSSD 明显高于前两条记录，提示其 RR 间期波动更大。",
        "呼吸率方面，bidmc_01、bidmc_02 和 bidmc_03 的平均呼吸率分别约为 17.15、12.45 和 14.76 次/min，均低于心率和脉率，符合基本生理规律。"
        "频谱和时频图也显示 RESP 主频位于低频区域，而 ECG 和 PPG 的主要周期性成分位于较高频率范围。这说明本文的滤波和特征提取结果在生理解释上具有合理性。",
        "Bland-Altman 分析显示，ECG 心率与 PPG 脉率整体偏差较小，但不同记录之间存在差异。PPG 是外周血容量变化信号，其峰值出现时间晚于 ECG R 峰，且容易受外周灌注和波形质量影响。"
        "因此，在记录 03 这类差异较明显的样本中，仅依赖 PPG 估计脉率可能带来一定误差。多模态融合分析的价值正在于通过 ECG、PPG 和 RESP 的互相验证，发现单一模态可能忽略的问题。",
        "本文仍存在局限性。第一，选取的真实记录数量为 3 条，虽然足以完成课程论文的完整处理流程，但尚不足以支持临床统计结论。第二，峰值检测采用基于距离和显著性的传统方法，"
        "在低质量波形上可能出现漏检或误检。第三，本文没有引入深度学习模型，主要原因是样本数量较少且课程目标更强调可解释的信号处理流程。未来可扩大记录数量，结合质量评估、异常窗口剔除和更稳健的峰值检测算法，进一步提高分析可靠性。"
    ]:
        add_paragraph(doc, text)

    doc.add_heading("9 结论", level=1)
    for text in [
        "本文基于 BIDMC PPG and Respiration Dataset，完成了 ECG、PPG 与呼吸信号的多模态生命体征分析。通过带通滤波、标准化、峰值检测、FFT 频谱分析和 STFT 时频分析，"
        "本文提取了心率、脉率、呼吸率、SDNN、RMSSD、RMS、主频率等特征，并使用多幅图表展示了原始波形、处理后波形、峰值检测、频谱、时频图和统计结果。",
        "实验结果表明，ECG 和 PPG 在多数记录中能够给出接近的心血管节律估计，呼吸信号则提供了独立的通气节律信息。跨记录比较、相关性矩阵、Bland-Altman 分析和 ECG-PPG 峰值时延估计进一步说明，多模态联合分析能够从电活动、外周血容量变化和呼吸节律三个角度综合描述生命体征。"
        "同时，实时可视化 GUI 的设计增强了项目的工程展示能力。总体来看，本文完成了数据获取、预处理、特征提取、可视化和结果讨论的完整流程，满足课程论文对生物医学信号公开数据集独立分析的要求。"
    ]:
        add_paragraph(doc, text)

    doc.add_heading("10 参考文献", level=1)
    refs = [
        "[1] Pimentel M A F, Johnson A E W, Charlton P H, et al. Towards a robust estimation of respiratory rate from pulse oximeters[J]. IEEE Transactions on Biomedical Engineering, 2017, 64(8): 1914-1923.",
        "[2] Goldberger A L, Amaral L A N, Glass L, et al. PhysioBank, PhysioToolkit, and PhysioNet: Components of a new research resource for complex physiologic signals[J]. Circulation, 2000, 101(23): e215-e220.",
        "[3] Task Force of the European Society of Cardiology and the North American Society of Pacing and Electrophysiology. Heart rate variability: standards of measurement, physiological interpretation and clinical use[J]. Circulation, 1996, 93(5): 1043-1065.",
        "[4] Elgendi M. On the analysis of fingertip photoplethysmogram signals[J]. Current Cardiology Reviews, 2012, 8(1): 14-25.",
        "[5] SciPy Developers. scipy.signal - Signal processing[EB/OL]. https://docs.scipy.org/doc/scipy/reference/signal.html.",
        "[6] Hunter J D. Matplotlib: A 2D graphics environment[J]. Computing in Science & Engineering, 2007, 9(3): 90-95.",
        "[7] PhysioNet. BIDMC PPG and Respiration Dataset[EB/OL]. https://physionet.org/content/bidmc/1.0.0/.",
    ]
    for ref in refs:
        add_paragraph(doc, ref, first_line=False)

    doc.add_heading("11 附录（代码说明）", level=1)
    for text in [
        "本文代码采用 Python 实现，主要依赖 NumPy、SciPy、Pandas 和 Matplotlib。项目目录包括 code、data、figures、results 和 docs 等文件夹。"
        "其中 data/raw 保存原始 BIDMC CSV 文件，data/processed 保存滤波和标准化后的数据，figures 保存论文图表，results 保存特征表和统计摘要。",
        "主要代码文件如下：main.py 负责单条记录的数据读取、预处理、特征提取和图表生成；aggregate_records.py 负责多条记录的统计汇总和跨记录比较；"
        "advanced_fusion_analysis.py 负责 ECG-PPG 一致性分析、峰值时延估计、心肺耦合和融合特征图；realtime_gui.py 负责实时可视化回放界面；download_bidmc_csv.py 提供数据下载辅助功能。",
        "运行顺序建议为：首先将 bidmc_01_Signals.csv、bidmc_02_Signals.csv 和 bidmc_03_Signals.csv 放入 data/raw；然后依次运行 python code/main.py --record 01、"
        "python code/main.py --record 02 和 python code/main.py --record 03；之后运行 python code/aggregate_records.py --records 01 02 03 和 python code/advanced_fusion_analysis.py --records 01 02 03。"
        "若需查看实时界面，可运行 python code/realtime_gui.py --record 01。"
    ]:
        add_paragraph(doc, text)


def main() -> None:
    root = project_root()
    out_dir = root / "docs"
    out_dir.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    add_title_page(doc)
    add_paper_body(doc, root)
    out = out_dir / "请填写学号_请填写姓名_课程论文.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
