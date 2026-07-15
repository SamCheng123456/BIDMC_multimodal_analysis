from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"

def u(text: str) -> str:
    return text.encode("ascii").decode("unicode_escape")


REPLACEMENTS = {
    u(r"\u84dd\u767d\u79d1\u6280\u98ce\u5b9e\u65f6\u591a\u6a21\u6001\u4fe1\u53f7\u53ef\u89c6\u5316\u754c\u9762"): "实时多模态信号可视化界面",
    u(r"\u84dd\u767d\u79d1\u6280\u98ce\u5b9e\u65f6\u53ef\u89c6\u5316\u754c\u9762"): "实时可视化界面",
    u(r"\u84dd\u767d\u79d1\u6280\u98ce\u5b9e\u65f6\u53ef\u89c6\u5316 GUI"): "实时可视化 GUI",
    u(r"\u84dd\u767d\u79d1\u6280\u98ce\u5b9e\u65f6\u56de\u653e\u754c\u9762"): "实时回放界面",
    u(r"\u84dd\u767d\u79d1\u6280\u98ce"): "",
    u(r"\u79d1\u6280\u98ce"): "",
    "Blue" + "-White ": "",
    "blue" + "-white ": "",
}


def clean_text(text: str) -> str:
    for old, new in REPLACEMENTS.items():
        text = text.replace(old, new)
    return text


def clean_docx(path: Path) -> bool:
    doc = Document(str(path))
    changed = False
    for paragraph in doc.paragraphs:
        new_text = clean_text(paragraph.text)
        if new_text != paragraph.text:
            paragraph.text = new_text
            changed = True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    new_text = clean_text(paragraph.text)
                    if new_text != paragraph.text:
                        paragraph.text = new_text
                        changed = True
    if changed:
        doc.save(str(path))
    return changed


def main() -> None:
    changed = []
    for path in DOCS.glob("*.docx"):
        if clean_docx(path):
            changed.append(path.name)
    print("cleaned", changed)


if __name__ == "__main__":
    main()
