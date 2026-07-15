from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
SOURCE = DOCS / "course_paper_full_dataset.docx"
OUTPUT = DOCS / "course_paper_full_dataset_enhanced.docx"


def insert_paragraph_before(doc: Document, startswith: str, text: str, style: str = "Normal"):
    target = next(p for p in doc.paragraphs if p.text.strip().startswith(startswith))
    paragraph = target.insert_paragraph_before(text)
    paragraph.style = style
    return paragraph


def insert_picture_before(doc: Document, startswith: str, image_path: Path, caption: str, width_cm: float = 15.2) -> None:
    target = next(p for p in doc.paragraphs if p.text.strip().startswith(startswith))
    cap = target.insert_paragraph_before(caption)
    cap.style = "Caption"
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para = target.insert_paragraph_before("")
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.add_run().add_picture(str(image_path), width=Cm(width_cm))


def replace_paragraph_containing(doc: Document, needle: str, text: str) -> None:
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            paragraph.text = text
            return
    raise ValueError(f"Paragraph not found: {needle}")


def main() -> None:
    doc = Document(str(SOURCE))

    insert_paragraph_before(
        doc,
        "3 数据预处理",
        "表 1 的意义不只是给出全数据集均值，还用于判断各类生命体征估计是否处于合理范围。"
        "从统计结果看，PPG 脉率与 ECG 心率在多数窗口中接近，说明两种心血管相关模态能够相互验证；"
        "但 HR-PR 绝对误差的均值明显高于中位数，提示误差分布并不对称，少数记录或少数窗口会显著拉高总体均值。"
        "SDNN 和 RMSSD 的标准差较大，说明不同记录之间心搏间期波动差异明显，也提示全数据集分析比少量示例记录更能暴露真实监护数据的复杂性。"
    )

    insert_picture_before(
        doc,
        "3 数据预处理",
        ROOT / "figures" / "full_dataset" / "fig0_overall_analysis_workflow.png",
        "总体流程图 ECG、PPG 与呼吸信号全数据集多模态分析流程",
    )

    replace_paragraph_containing(
        doc,
        "除离线分析外，本文设计了一个",
        "除离线分析外，本文设计了一个实时可视化界面，用于回放处理后的 ECG、PPG 和 RESP 三路信号。"
        "界面基于 Python 标准库 Tkinter 和 Matplotlib 实现，不需要额外复杂框架。"
        "用户可以选择记录编号，设置信号显示窗口长度和播放步长，通过播放、暂停、重置和滑动条查看不同时间段的多模态信号。"
        "GUI 的功能重点并不是单纯展示静态图片，而是支持记录选择、窗口长度设置、播放步长设置、暂停与重置、时间轴拖动和三路波形同步查看，"
        "使离线计算得到的多模态信号能够以接近真实软件的方式进行交互式复核。界面中的波形绘制和图表导出主要使用 Matplotlib 完成[6]。"
    )

    replace_paragraph_containing(
        doc,
        "新界面整体以浅蓝背景和白色内容卡片为基础",
        "界面整体采用简洁的监护面板布局。顶部状态栏显示项目名称、记录编号和当前时间窗口；"
        "中部设置四个指标卡片，用于展示 ECG、PPG、RESP 三路标准化信号的动态范围和采样率；"
        "下方采用三个同步波形区分别显示三路信号，并通过颜色和标签区分模态。界面底部提供播放、重置和时间轴滑动控制。"
    )

    replace_paragraph_containing(
        doc,
        "\u56fe 17 \u84dd\u767d\u79d1\u6280\u98ce",
        "图 17 实时多模态信号可视化界面",
    )

    insert_picture_before(
        doc,
        "7.1 全数据集扩展结果",
        ROOT / "figures" / "gui" / "fig_gui_software_explainer.png",
        "GUI 说明图 实时回放界面的功能结构",
    )

    insert_paragraph_before(
        doc,
        "图 13 ECG-PPG 峰值时延分布",
        "表 2 进一步说明，多模态融合结果需要与质量控制结果结合解读。"
        "ECG-PPG 峰值时延样本数反映了两种心血管模态在时间轴上可匹配的搏动事件数量；"
        "PAT 均值用于描述 ECG 电活动到 PPG 外周脉搏波之间的时间耦合，但本文不将其直接解释为严格临床脉搏传导时间。"
        "质量控制通过记录与需复核记录的划分表明，全数据集中的大多数窗口与监护仪 Numerics 一致，但部分记录仍可能存在 ECG 过检、PPG 波形质量下降或呼吸信号异常。"
        "因此，表 2 的价值在于把“结果估计”和“结果可信度”放在同一框架下讨论。"
    )

    insert_paragraph_before(
        doc,
        "9 结论",
        "关于深度学习方法，本文没有将其作为主要实验路线。原因在于，本研究的核心目标是完成可解释的生物医学信号处理流程，"
        "包括滤波、峰值检测、频域分析、时频分析、HRV 特征、多模态一致性和质量控制。"
        "若在缺少逐搏人工标注和严格训练/测试划分的情况下直接加入深度学习模型，容易造成模型目标不清或数据泄漏，"
        "反而削弱结果可信度。本文将深度学习保留为未来扩展方向：可在现有 848 个窗口级样本基础上，"
        "以 Numerics 误差阈值构造质量复核标签，训练轻量级 MLP 或 1D-CNN 对异常窗口进行辅助筛选。"
        "但这类模型只能作为探索性质量控制工具，不能替代本文基于生理机制和信号处理方法得到的主要结论。"
    )

    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
