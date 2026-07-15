from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
SOURCE = DOCS / "course_paper_format_checked_cited_fixed.docx"
OUTPUT = DOCS / "course_paper_full_dataset.docx"


def fmt(value: float, digits: int = 2) -> str:
    return f"{value:.{digits}f}"


def set_cell_text(cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "宋体"
            run.font.size = Pt(9)


def replace_table_after_caption(doc: Document, caption_prefix: str, headers: list[str], rows: list[list[str]]) -> None:
    caption = next(p for p in doc.paragraphs if p.text.strip().startswith(caption_prefix))
    if caption_prefix.startswith("表 1"):
        old_table = doc.tables[0]
    else:
        old_table = doc.tables[1 if len(doc.tables) > 1 else 0]
    old_table._element.getparent().remove(old_table._element)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_text(cell, text)
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)
    caption._p.addnext(table._tbl)


def replace_paragraph_containing(doc: Document, needle: str, text: str) -> None:
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            paragraph.text = text
            return
    raise ValueError(f"Paragraph not found: {needle}")


def insert_full_dataset_figures(doc: Document) -> None:
    target = next(p for p in doc.paragraphs if p.text.strip().startswith("8 分析与讨论"))
    heading = target.insert_paragraph_before("7.1 全数据集扩展结果")
    heading.style = "Heading 2"
    intro = target.insert_paragraph_before(
        "在完成单条记录的波形展示和多模态特征提取后，本文进一步对 BIDMC 全部 53 条记录进行批处理分析。"
        "全量分析共得到 848 个 30 s 窗口，并结合 Numerics 文件中的监护仪 HR、PULSE 和 RESP 数值进行质量控制对比。"
        "以下图表用于展示全数据集层面的生命体征分布、模态一致性和需要复核的记录。"
    )
    intro.style = "Normal"

    figures = [
        ("fig19_full_dataset_vital_sign_distributions.png", "图 18 全数据集生命体征分布"),
        ("fig20_ranked_hr_pr_error_all_records.png", "图 19 全部记录 ECG 心率与 PPG 脉率误差排序"),
        ("fig21_full_dataset_feature_correlation.png", "图 20 全数据集多模态特征相关性矩阵"),
        ("fig23_signal_vs_monitor_error_distribution.png", "图 21 信号估计值与监护仪 Numerics 的误差分布"),
        ("fig24_signal_monitor_agreement_scatter.png", "图 22 信号估计值与监护仪 Numerics 的一致性散点图"),
        ("fig25_record_quality_flags.png", "图 23 记录级 ECG 与监护仪 HR 误差排序及质量标记"),
    ]
    for filename, caption in figures:
        cap = target.insert_paragraph_before(caption)
        cap.style = "Caption"
        para = target.insert_paragraph_before("")
        run = para.add_run()
        run.add_picture(str(ROOT / "figures" / "full_dataset" / filename), width=Cm(15.2))


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def remove_caption_and_previous_image(doc: Document, caption_prefixes: tuple[str, ...]) -> None:
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if text.startswith(caption_prefixes):
            previous = paragraph._p.getprevious()
            if previous is not None:
                previous.getparent().remove(previous)
            remove_paragraph(paragraph)


def main() -> None:
    summary = pd.read_csv(ROOT / "results" / "full_dataset_record_summary.csv", dtype={"record_id": str})
    features = pd.read_csv(ROOT / "results" / "full_dataset_window_features.csv", dtype={"record_id": str})
    monitor = pd.read_csv(ROOT / "results" / "full_dataset_monitor_comparison_summary.csv", dtype={"record_id": str})
    fusion = pd.read_csv(ROOT / "results" / "advanced_fusion_summary.csv", dtype={"record_id": str})

    doc = Document(str(SOURCE))

    replace_paragraph_containing(
        doc,
        "选取 01、02、03 三条代表性记录",
        "为提高对生命体征状态的综合理解，本文基于 PhysioNet 平台公开的 BIDMC PPG and Respiration Dataset，"
        "对全部 53 条记录中同步采集的 ECG、PPG 与呼吸信号进行多模态分析。研究首先修正 CSV 时间列三位小数取整带来的采样率估计偏差，"
        "按 125 Hz 采样率对三路信号进行带通滤波和标准化处理，其中 ECG 采用 0.5-40 Hz 带通滤波，PPG 采用 0.5-8 Hz 带通滤波，"
        "呼吸信号采用 0.05-1 Hz 带通滤波。随后从时域、频域和时频域三个角度开展分析，提取心率、脉率、呼吸率、SDNN、RMSSD、RMS、"
        "主频率等特征，并进一步构建 ECG-PPG 一致性分析、ECG-PPG 峰值时延估计、心肺耦合分析以及与监护仪 Numerics 的质量控制对比。"
        "全数据集共得到 848 个 30 s 分析窗口。结果显示，大多数窗口的信号估计值与监护仪数值具有较好一致性，窗口级中位绝对误差分别为 "
        "ECG 心率 0.66 bpm、PPG 脉率 0.65 bpm 和呼吸率 0.51 次/min；同时有 13 条记录被标记为需要复核，说明全数据集分析必须结合信号质量控制。"
    )

    replace_paragraph_containing(
        doc,
        "本文选取 bidmc_01、bidmc_02 和 bidmc_03",
        "本文使用的数据集为 PhysioNet 平台发布的 BIDMC PPG and Respiration Dataset v1.0.0。"
        "该数据集来源于 Beth Israel Deaconess Medical Center 的危重患者监护记录，包含 ECG、PPG/PLETH、阻抗呼吸信号、生命体征数值以及呼吸相关标注。"
        "数据集总共包含 53 条记录，每条记录约 8 min，波形信号采样率为 125 Hz。本文不再只选取 3 条课程复现实验子集，而是处理 bidmc_01 至 bidmc_53 的全部 Signals 文件，"
        "并读取对应 Numerics 文件作为 HR、PULSE 和 RESP 的监护仪参考数值。每条记录按 30 s 窗口切分，全部记录共得到 848 个分析窗口。"
        "上述数据集说明与记录结构以 PhysioNet 官方数据集页面为依据[7]；该数据也曾用于从脉搏血氧波形估计呼吸率的相关研究[1]。"
    )

    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("表 1"):
            paragraph.text = "表 1 全部 53 条 BIDMC 记录的主要生命体征与 HRV 特征统计"
        if paragraph.text.strip().startswith("表 2"):
            paragraph.text = "表 2 全数据集多模态融合与质量控制摘要"

    table1_rows = []
    table_metrics = [
        ("ECG心率(bpm)", "ecg_hr_bpm"),
        ("PPG脉率(bpm)", "ppg_pr_bpm"),
        ("呼吸率(次/min)", "resp_rr_bpm"),
        ("SDNN(ms)", "ecg_sdnn_ms"),
        ("RMSSD(ms)", "ecg_rmssd_ms"),
        ("HR-PR绝对误差(bpm)", "hr_pr_abs_error_bpm"),
    ]
    for label, col in table_metrics:
        s = features[col]
        table1_rows.append([label, f"{fmt(s.mean())}±{fmt(s.std())}", fmt(s.median()), fmt(s.min()), fmt(s.max())])
    replace_table_after_caption(doc, "表 1", ["指标", "均值±标准差", "中位数", "最小值", "最大值"], table1_rows)

    ok_count = int((monitor["signal_quality_flag"] == "ok").sum())
    review_count = int((monitor["signal_quality_flag"] == "review").sum())
    table2_rows = [
        ["分析记录数", "53", "bidmc_01 至 bidmc_53 全部记录"],
        ["分析窗口数", str(len(features)), "30 s 非重叠窗口"],
        ["ECG-PPG峰值时延样本数", str(int(fusion["pat_count"].sum())), "用于时间耦合分析"],
        ["PAT记录均值(ms)", fmt(fusion["pat_mean_ms"].mean()), "对各记录 PAT 均值再求平均"],
        ["质量控制通过记录", str(ok_count), "与 Numerics 对比误差较小"],
        ["需复核记录", str(review_count), "至少一个模态误差超过阈值"],
        ["窗口级ECG误差中位数(bpm)", "0.66", "ECG HR 对监护仪 HR"],
        ["窗口级PPG误差中位数(bpm)", "0.65", "PPG PR 对监护仪 PULSE"],
        ["窗口级RESP误差中位数(次/min)", "0.51", "RESP RR 对监护仪 RESP"],
    ]
    replace_table_after_caption(doc, "表 2", ["项目", "数值", "说明"], table2_rows)

    replace_paragraph_containing(
        doc,
        "从三条记录的结果看",
        "从全部 53 条记录的结果看，ECG 心率、PPG 脉率和呼吸率的总体分布均处于临床监护数据可解释范围内。"
        f"窗口级 ECG 心率均值为 {fmt(features['ecg_hr_bpm'].mean())} bpm，PPG 脉率均值为 {fmt(features['ppg_pr_bpm'].mean())} bpm，"
        f"呼吸率均值为 {fmt(features['resp_rr_bpm'].mean())} 次/min。与只分析少数记录相比，全量分析更能暴露真实数据中的质量差异："
        "多数记录的 ECG 与 PPG 估计结果接近，但少数记录存在明显偏差，尤其是 ECG 峰值过检时会导致心率估计约为监护仪 HR 的两倍。"
    )

    replace_paragraph_containing(
        doc,
        "图 9 展示了记录 01 的分窗生命体征趋势",
        "图 9 保留记录 01 的分窗生命体征趋势作为单条记录示例，用于说明窗口级心率、脉率和呼吸率的计算方式。"
        "在全数据集层面，图 18 至图 23 进一步展示 53 条记录的生命体征分布、ECG-PPG 误差排序、多模态特征相关性以及与监护仪 Numerics 的质量控制对比。"
    )

    replace_paragraph_containing(
        doc,
        "图 11 的 ECG 心率与 PPG 脉率误差图是多模态融合的重要结果",
        "全数据集结果显示，ECG 与 PPG 在多数窗口具有较好一致性，但记录间差异明显。"
        "图 19 的误差排序能够快速定位 ECG-PPG 差异较大的记录，图 21 和图 22 则从监护仪 Numerics 角度验证信号估计结果。"
        "这些结果说明，多模态融合不应只追求平均误差较小，还应识别信号质量较差或峰值检测不稳定的记录。"
    )

    replace_paragraph_containing(
        doc,
        "本文仍存在局限性。第一，选取的真实记录数量为 3 条",
        "本文仍存在局限性。第一，虽然本文已处理全部 53 条记录，但峰值检测仍主要采用传统规则方法，在部分低质量 ECG 或 PPG 波形上可能出现漏检或过检。"
        "第二，Numerics 文件提供的是监护仪数值，并不等同于逐搏人工标注，因此本文将其作为质量控制参照，而不是严格逐点真值。"
        "第三，本文没有引入深度学习模型，主要原因是课程目标更强调可解释的信号处理流程，且深度学习需要更细致的标注、训练集划分和泛化验证。"
        "未来可在全量记录基础上增加信号质量指数、异常窗口剔除和更稳健的峰值检测算法。"
    )

    replace_paragraph_containing(
        doc,
        "本文基于 BIDMC PPG and Respiration Dataset，完成了 ECG、PPG 与呼吸信号的多模态生命体征分析。",
        "本文基于 BIDMC PPG and Respiration Dataset，完成了全部 53 条记录的 ECG、PPG 与呼吸信号多模态生命体征分析。"
        "通过带通滤波、标准化、峰值检测、FFT 频谱分析和 STFT 时频分析，本文提取了心率、脉率、呼吸率、SDNN、RMSSD、RMS、主频率等特征，"
        "并使用全数据集图表展示了生命体征分布、模态一致性、特征相关性和质量控制结果。"
    )

    replace_paragraph_containing(
        doc,
        "实验结果表明，ECG 和 PPG 在多数记录中能够给出接近的心血管节律估计",
        "实验结果表明，ECG 和 PPG 在大多数窗口中能够给出接近的心血管节律估计，呼吸信号则提供了独立的通气节律信息。"
        "与 Numerics 的对比显示，窗口级 ECG 心率、PPG 脉率和呼吸率的中位绝对误差均较小，但仍有 13 条记录需要复核。"
        "这说明多模态联合分析不仅能够估计生命体征，也能够发现单一信号或单一算法难以暴露的质量问题。"
        "同时，实时可视化 GUI 的设计增强了项目的工程展示能力。总体来看，本文完成了全数据集数据获取、预处理、特征提取、可视化、质量控制和结果讨论的完整流程。"
    )

    remove_caption_and_previous_image(doc, ("图 10 三条记录", "图 11 ECG 心率", "图 12 跨记录"))
    insert_full_dataset_figures(doc)
    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
