from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def project_root() -> Path:
    return Path(__file__).resolve().parents[1]


def set_cn(run, font: str = "宋体", size: int = 11, bold: bool = False) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold


def add_para_before(anchor, text: str, first_line: bool = True):
    p = anchor.insert_paragraph_before()
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(6)
    if first_line:
        p.paragraph_format.first_line_indent = Pt(22)
    run = p.add_run(text)
    set_cn(run)
    return p


def add_heading_before(anchor, text: str, level: int = 2):
    p = anchor.insert_paragraph_before()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    set_cn(run, "黑体", 14 if level == 2 else 12, True)
    return p


def add_figure_before(anchor, path: Path, caption: str, width_cm: float = 14.0):
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = anchor.insert_paragraph_before()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_cn(r, "宋体", 9)


def find_para(doc: Document, text: str):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise ValueError(f"Cannot find paragraph: {text}")


def main() -> None:
    root = project_root()
    docs_dir = root / "docs"
    source = docs_dir / "请填写学号_请填写姓名_课程论文.docx"
    if not source.exists():
        candidates = sorted(docs_dir.glob("*课程论文.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
        if not candidates:
            raise FileNotFoundError("No paper docx found.")
        source = candidates[0]

    doc = Document(str(source))
    gui_fig = root / "figures" / "gui" / "fig18_gui_preview.png"

    anchor_8 = find_para(doc, "8 分析与讨论")
    add_heading_before(anchor_8, "7.1 实时可视化界面展示", level=2)
    add_para_before(
        anchor_8,
        "根据课程任务书中“实时可视化界面 / GUI 设计”的鼓励方向，本文对原有回放工具进行了界面重构。新界面"
        "整体以浅蓝背景和白色内容卡片为基础，顶部使用深蓝标题栏显示项目名称、记录编号和当前时间窗口，既保证视觉层次，也使界面风格与生物医学监护场景相匹配。"
        "中部设置四个指标卡片，用于展示 ECG、PPG、RESP 三路标准化信号的动态范围和采样率；下方采用三个同步波形区分别显示三路信号，并使用不同蓝色系进行区分。"
    )
    add_para_before(
        anchor_8,
        "从界面展示图可以看出，ECG 波形具有尖锐而密集的 R 峰，PPG 波形具有较平滑的脉搏上升支和下降支，RESP 波形则呈现明显的低频周期起伏。"
        "这种同屏同步显示方式比单独查看静态图更直观，能够帮助观察三种模态在同一时间轴上的相对变化。例如，当 ECG 出现一次 R 峰后，PPG 波形通常会在短暂延迟后出现脉搏峰，"
        "而呼吸信号的变化周期明显更长。该界面因此不仅是展示工具，也能辅助理解 ECG-PPG 时延估计和心肺耦合分析的生理含义。"
    )
    add_para_before(
        anchor_8,
        "界面底部提供播放、重置和时间轴滑动控制，用户可以连续回放信号，也可以拖动到某一时间段进行局部观察。"
        "与最终论文中的静态图表相比，GUI 更接近实际监护系统中的观察方式，体现了本项目从离线分析到交互展示的扩展。"
        "考虑到课程论文的重点仍是生物医学信号处理，GUI 没有引入复杂框架，而是使用 Tkinter 和 Matplotlib 实现，保证代码简洁、可运行、易复现。"
    )
    add_figure_before(anchor_8, gui_fig, "图 17 实时多模态信号可视化界面", width_cm=15.2)

    anchor_8 = find_para(doc, "8 分析与讨论")
    add_heading_before(anchor_8, "7.2 图表结果的进一步分析", level=2)
    detailed = [
        "图 1 展示了记录 01 中原始 ECG、PPG 和 RESP 信号的整体形态。原始 ECG 中可以观察到幅值较高、持续时间较短的尖峰，这些尖峰对应心室去极化形成的 R 峰，"
        "是心搏定位和 RR 间期计算的主要依据。PPG 原始波形虽然也具有周期性，但其峰形比 ECG 更宽，且上升支和下降支不完全对称，说明 PPG 反映的是外周血容量变化，"
        "并非心脏电活动本身。RESP 信号变化最慢，一个呼吸周期持续数秒，说明其主频明显低于 ECG 和 PPG。三者在同一时间轴上的差异，正是本文进行多模态联合分析的基础。",
        "图 2 与图 1 对比后可以看出，滤波后 ECG 的基线起伏被压低，R 峰轮廓更加清晰；PPG 波形中部分高频毛刺被削弱，脉搏波周期更容易辨认；RESP 信号经过低频带通处理后，"
        "呼吸周期的上升和下降趋势更加平滑。该结果说明所选滤波参数与三类信号的生理频段基本匹配。特别是呼吸信号，如果使用过宽频带或普通滤波形式，容易引入噪声或数值不稳定，"
        "因此本文使用二阶节零相位滤波，使低频处理更稳定。",
        "图 3 将三路信号标准化后叠加显示。标准化后的曲线失去了原始幅值单位，但保留了时间位置和相对波形变化，因此适合观察不同模态之间的同步关系。"
        "在局部时间段内，ECG R 峰出现后，PPG 脉搏峰并非同时出现，而是存在一定延迟；RESP 则以更长周期缓慢变化。这种现象说明 ECG、PPG 和 RESP 分别对应不同生理过程，"
        "直接比较原始幅值意义不大，而比较时间关系和周期结构更有价值。",
        "图 4、图 5 和图 6 分别展示了三类峰值检测结果。ECG R 峰检测相对稳定，因为 R 峰具有高幅值和尖锐形态；PPG 峰值检测更依赖波形质量，局部波形变宽或幅值变化时，"
        "峰值位置可能产生偏移；RESP 峰值数量远少于 ECG 和 PPG，符合呼吸频率低于心率的基本规律。三类峰值检测图共同证明，本文并不是只做整体统计，而是先在信号层面定位具体生理事件，"
        "再基于事件间期计算生命体征特征。",
        "图 7 的频谱结果从频域角度验证了时域观察。ECG 与 PPG 的主要能量位于心率相关频段，RESP 的主要能量则集中在更低频率范围。"
        "如果某一记录中 PPG 频谱主峰与 ECG 主峰偏离较大，则可能提示 PPG 波形质量下降或峰值检测存在误差。图 8 的 STFT 时频图进一步表明，真实监护信号具有非平稳性，"
        "即频率成分会随时间轻微变化，因此仅用全局平均频谱描述信号是不充分的。",
        "图 9 展示了记录 01 的分窗生命体征趋势。可以看到 ECG 心率和 PPG 脉率在多数窗口中走势接近，说明二者能够相互验证；呼吸率曲线整体位于更低数值范围，且变化幅度较小。"
        "图 10 将三条记录进行横向比较，发现记录 01 和 02 的心率、脉率均接近 73 bpm，而记录 03 的心率和脉率明显较低。"
        "这说明不同患者或不同监护片段之间存在个体差异，跨记录比较比单条记录分析更能体现数据集的真实变化。",
        "图 11 的 ECG 心率与 PPG 脉率误差图是多模态融合的重要结果。记录 01 和 02 的平均绝对误差小于 1 bpm，说明 PPG 在这两条记录中可以较好估计脉率；"
        "记录 03 的误差增大到 3.40 bpm，提示 PPG 结果不能无条件替代 ECG。图 12 的相关性矩阵进一步展示多特征之间的统计关系，可用于发现哪些指标变化方向一致，"
        "哪些指标具有相对独立的信息。",
        "图 13 至图 16 是扩展融合分析结果。ECG-PPG 峰值时延分布反映心电活动和外周脉搏波之间的时间关系，Bland-Altman 图则直接评估 ECG 心率和 PPG 脉率的一致性。"
        "从图 14 可以看出，记录 01 和 02 的点大多靠近零差值附近，而记录 03 出现更明显偏差，这与图 11 的误差统计一致。图 15 的心肺耦合散点图用于观察呼吸率与心率之间的关系，"
        "图 16 则将多个融合指标归一化后集中展示，使跨记录差异更加直观。"
    ]
    for text in detailed:
        add_para_before(anchor_8, text)

    anchor_9 = find_para(doc, "9 结论")
    add_heading_before(anchor_9, "8.1 关于结果可靠性的补充讨论", level=2)
    for text in [
        "从结果可靠性角度看，本文的分析具有两个优点。第一，生命体征估计不是只依赖一种方法，而是同时由峰值检测、频域主峰和分窗趋势进行交叉验证。"
        "例如，ECG 与 PPG 在记录 01 和 02 中得到接近的心率或脉率，说明两种模态在这些片段中具有较好一致性。第二，本文保留了跨记录差异，而不是强行把所有记录解释成同一种趋势。"
        "记录 03 的 HR-PR 误差较大，恰好说明真实数据存在质量差异，也说明多模态分析能够发现单一信号估计中的潜在风险。",
        "需要强调的是，本文并未将 ECG-PPG 峰值时延直接解释为严格临床意义上的脉搏传导时间。临床 PTT 通常需要明确传感器位置、血管路径和同步采集条件，而本文使用的是公开监护数据中的 ECG 与 PPG 波形，"
        "因此更稳妥的表述是“峰值时延估计”或“时间耦合特征”。这种表述既保留了创新性，也避免过度医学化解释。",
        "相比加入深度学习模型，本文选择增强可解释的信号处理和多模态融合分析。原因在于本研究只选取 3 条记录作为课程论文样本，样本量不足以训练可靠神经网络。"
        "在这种条件下，盲目加入深度学习容易造成模型结果缺乏统计意义。本文采用滤波、峰值检测、HRV 特征、一致性分析和 GUI 展示，更符合课程对独立分析、代码可复现和结果可解释的要求。"
    ]:
        add_para_before(anchor_9, text)

    out = docs_dir / "请填写学号_请填写姓名_课程论文_增强版.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
